from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
LAB_DIR = ROOT / "labs"
OUT = ROOT / "Assessments"
ASSETS = OUT / "assets"

COURSE_TITLE = "Microsoft Certified Azure AI Engineer Associate AI-102 Training"
COURSE_CODE = "TGS-2023036651"
ORG = "Tertiary Infotech Academy Pte Ltd"
UEN = "201200696W"
VERSION = "v1.0"

BLUE = "1F6FEB"
DARK = "111827"
MUTED = "4B5563"
LIGHT = "F3F7FC"
RULE = "D9E2F0"


ANSWER_KEY = {
    "Why is service selection part of AI engineering?": "The engineer maps business requirements to the right Azure AI services, deployment model, integrations, constraints, cost, security and operational controls so the solution is feasible and maintainable.",
    "Why should keys be protected?": "Keys grant access to Azure AI resources. They must be stored securely, rotated, and replaced with managed identity or least-privilege access where possible.",
    "What should be monitored in an AI service?": "Monitor usage, latency, errors, token or transaction consumption, cost, model/output quality signals, safety events, authentication failures and availability.",
    "What is the role of responsible AI planning?": "It identifies risks early and defines fairness, safety, privacy, transparency, accountability, escalation and monitoring controls before production deployment.",
    "What are content filters?": "Controls that detect or block harmful categories of input or output, such as violence, hate, sexual content, self-harm, jailbreak attempts or unsafe generated responses.",
    "What is a prompt shield?": "A control that helps detect and mitigate prompt injection, jailbreak and indirect prompt attacks that try to override system instructions or expose data.",
    "Why is human escalation important?": "It provides accountability and review for uncertain, harmful, high-impact or low-confidence cases where automated output should not be final.",
    "Who is accountable for AI output?": "The organization deploying the system remains accountable through its product owner, governance process and assigned responsible personnel.",
    "What is RAG?": "Retrieval-augmented generation combines search or retrieval from trusted sources with a generative model so answers are grounded in approved content.",
    "Why are prompt templates useful?": "They standardize instructions, input variables, constraints, tone and output format so model behavior is more repeatable and easier to evaluate.",
    "What does tracing help debug?": "Tracing helps inspect prompt flow execution, retrieved context, tool calls, intermediate outputs, failures, latency and data passed between steps.",
    "How can feedback improve a generative AI solution?": "Feedback supplies examples for evaluation, prompt improvement, grounding refinements, safety tuning and prioritizing recurring failure cases.",
    "What is an AI agent?": "An AI system that can reason over a task, use tools or knowledge sources, follow instructions and take controlled actions toward a goal.",
    "Why do agents need tool permissions?": "Permissions limit what an agent can access or change, enforcing least privilege and reducing the impact of erroneous or malicious actions.",
    "What is orchestration?": "The coordination of model reasoning, tools, workflows, memory, retrieval, handoffs and sometimes multiple agents to complete a task safely.",
    "Why should autonomous actions be constrained?": "Constraints prevent uncontrolled decisions, accidental changes, privacy leaks, cost spikes, unsafe actions and actions without human approval.",
    "What is OCR?": "Optical character recognition extracts machine-readable text from images, screenshots, scanned forms or documents.",
    "How is object detection different from classification?": "Classification assigns labels to an image, while object detection identifies objects and their locations, often with bounding boxes.",
    "What does a custom vision model require?": "A defined task, representative labeled images, training and validation data, a selected model type, evaluation metrics and ongoing review for drift and bias.",
    "Why is video analysis sensitive?": "Video may include faces, voices, locations, behavior or personal data, so privacy, consent, retention, security and bias controls are required.",
    "What is entity recognition?": "A language capability that identifies named entities such as people, locations, organizations, dates, quantities and domain-specific terms in text.",
    "Why detect PII before storing text?": "PII detection supports privacy protection, masking, minimization, compliance and safer downstream analytics or model use.",
    "What is SSML?": "Speech Synthesis Markup Language controls text-to-speech output such as voice, pauses, pronunciation, emphasis, rate and pitch.",
    "When should translation outputs be reviewed?": "Review is needed for regulated, customer-facing, ambiguous, sensitive, legal, medical, safety or high-impact communications.",
    "What is an intent?": "The user goal or action expressed in an utterance, such as checking order status or requesting a refund.",
    "What is an entity?": "A structured value extracted from user language, such as product name, order number, date, location or support category.",
    "Why add alternate phrasing?": "Alternate utterances improve recognition of the same intent across varied wording, spelling, tone and language patterns.",
    "Why should language projects be backed up?": "Backups protect training data, labels, configuration, versions and recovery plans if a project is corrupted, deleted or changed incorrectly.",
    "What is an indexer?": "A component in Azure AI Search that reads content from a data source, extracts fields, applies enrichment if configured and populates an index.",
    "What is a skillset?": "A pipeline of enrichment skills such as OCR, entity recognition, key phrase extraction, translation or custom skills applied during indexing.",
    "What is vector search?": "Search that compares embeddings to find semantically similar content, even when the exact keywords differ.",
    "Why use semantic ranking?": "Semantic ranking improves relevance by re-ranking results using language understanding rather than relying only on keyword matching.",
    "When should you use a prebuilt model?": "Use a prebuilt model when the document type is supported and the required fields match common forms such as invoices, receipts, IDs or general documents.",
    "What is a composed model?": "A composed Document Intelligence model routes documents to one of several custom models so different form types can be processed through a single endpoint.",
    "Why are confidence scores useful?": "They help decide whether extraction can be accepted automatically, needs review or should trigger exception handling.",
    "What content types can Content Understanding process?": "It can process multiple modalities such as documents, images, audio, video and text depending on the configured analyzer and scenario.",
    "Which service should power grounded document search?": "Azure AI Search should power document retrieval and grounding, often combined with Azure OpenAI or Foundry for RAG.",
    "When should a human review AI output?": "Human review is needed for low-confidence, high-impact, harmful, sensitive, regulated, ambiguous or escalated outputs.",
    "What is the difference between RAG and fine-tuning?": "RAG retrieves external knowledge at runtime to ground answers; fine-tuning changes model behavior using training examples and does not itself provide current source knowledge.",
    "Which AI workload do you need to review most?": "The learner should identify their weakest AI-102 workload and justify it using confidence ratings, lab evidence and a practical review plan.",
}


def read_block(text, heading):
    match = re.search(rf"^## {re.escape(heading)}\s*\n(.*?)(?=^## |\Z)", text, re.S | re.M)
    return match.group(1).strip() if match else ""


def bullets(block):
    return [line.strip()[2:] for line in block.splitlines() if line.strip().startswith("- ")]


def numbered(block):
    values = []
    for line in block.splitlines():
        match = re.match(r"\s*\d+\.\s+(.*)", line)
        if match:
            values.append(match.group(1).strip())
    return values


def load_labs():
    labs = []
    for path in sorted(LAB_DIR.glob("lab-*.md")):
        text = path.read_text(encoding="utf-8")
        title = re.search(r"^#\s+(.+)", text, re.M).group(1)
        lab_no = int(re.search(r"Lab\s+(\d+)", title).group(1))
        slide_start = 17 + (lab_no - 1) * 5
        if lab_no >= 7:
            slide_start += 1
        labs.append(
            {
                "no": lab_no,
                "code": f"LO{lab_no}",
                "title": title,
                "short": title.split(" - ", 1)[1],
                "source": f"PPT slides {slide_start}-{slide_start + 4}; Lab {lab_no}",
                "slides": f"{slide_start}-{slide_start + 4}",
                "scenario": read_block(text, "Scenario"),
                "objectives": bullets(read_block(text, "Objectives")),
                "steps": [m.group(1).strip() for m in re.finditer(r"^###\s+(.+)", text, re.M)],
                "validation": read_block(text, "Validation").strip(),
                "questions": numbered(read_block(text, "Checkpoint Questions")),
            }
        )
    return labs


def pp_evidence(lab):
    evidence = {
        1: ["Service-to-requirement map.", "Foundry/resource plan.", "Security controls.", "Monitoring and cost plan.", "Architecture diagram."],
        2: ["Responsible AI risk register.", "Content safety control matrix.", "Safety policy.", "Governance scenario decisions.", "Incident response plan."],
        3: ["Model selection notes.", "Prompt templates.", "Prompt flow diagram.", "RAG design.", "Evaluation table."],
        4: ["Agent use case.", "Resource plan.", "Tool permission table.", "Orchestration diagram.", "Multi-agent review notes."],
        5: ["Vision requirement map.", "Image analysis request plan.", "Custom vision training plan.", "Video insight plan.", "Responsible AI review."],
        6: ["Language task map.", "Text analysis pipeline.", "Speech pipeline.", "SSML notes.", "Translation review plan."],
        7: ["Intent and entity table.", "Training data plan.", "Question answering plan.", "Multilingual strategy.", "Backup and recovery notes."],
        8: ["Index design.", "Data source and indexer plan.", "Skillset design.", "Query examples.", "Knowledge store projection notes."],
        9: ["Document scenario map.", "Custom model plan.", "Content understanding plan.", "Human review rules.", "Integration flow."],
        10: ["End-to-end service map.", "Architecture diagram.", "Governance checklist.", "Confidence matrix.", "Cleanup and 7-day review plan."],
    }
    return evidence[lab["no"]]


def ensure_assets():
    OUT.mkdir(parents=True, exist_ok=True)
    ASSETS.mkdir(parents=True, exist_ok=True)
    try:
        from PIL import Image, ImageDraw, ImageFont
    except Exception:
        return None, None

    logo_path = ASSETS / "tertiary-infotech-logo.png"
    badge_path = ASSETS / "ai102-course-badge.png"
    if not logo_path.exists():
        img = Image.new("RGBA", (360, 360), (255, 255, 255, 0))
        d = ImageDraw.Draw(img)
        d.ellipse((20, 20, 340, 340), fill=(13, 63, 148, 255))
        font = ImageFont.truetype("arialbd.ttf", 210)
        d.text((180, 176), "T", fill="white", anchor="mm", font=font)
        img.save(logo_path)
    if not badge_path.exists():
        img = Image.new("RGBA", (360, 360), (255, 255, 255, 0))
        d = ImageDraw.Draw(img)
        d.ellipse((20, 20, 340, 340), fill=(36, 114, 231, 255), outline=(214, 226, 240, 255), width=10)
        d.rounded_rectangle((72, 155, 288, 220), radius=18, fill=(16, 185, 129, 255))
        small = ImageFont.truetype("arial.ttf", 32)
        med = ImageFont.truetype("arialbd.ttf", 68)
        tiny = ImageFont.truetype("arialbd.ttf", 22)
        d.text((180, 95), "Microsoft Azure", fill="white", anchor="mm", font=small)
        d.text((180, 188), "AI-102", fill="white", anchor="mm", font=med)
        d.text((180, 248), "AI Engineer", fill="white", anchor="mm", font=tiny)
        img.save(badge_path)
    return logo_path, badge_path


def set_font(run, size=11, bold=False, color=DARK, italic=False):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_text(cell, text, bold=False, color=DARK, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def para(doc, text="", size=11, bold=False, color=DARK, italic=False, align=None, before=0, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, color=color, italic=italic)
    return p


def add_bullets(doc, items, style="List Bullet", size=10.5):
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        set_font(r, size=size)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    set_font(run, size=8, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    page_text = OxmlElement("w:t")
    page_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(page_text)
    run._r.append(fld_end)


def add_footer(doc):
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = ""
    r = footer.add_run(f"(c) 2026 {ORG}. All rights reserved.  |  www.tertiarycourses.com.sg")
    set_font(r, size=8, color=MUTED)
    r = footer.add_run("  |  Page ")
    set_font(r, size=8, color=MUTED)
    add_page_number_field(footer)


def base_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    add_footer(doc)
    return doc


def cover_page(doc, instrument, org_logo=None, course_logo=None):
    for _ in range(2):
        para(doc, "", after=10)
    if org_logo:
        p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
        p.add_run().add_picture(str(org_logo), width=Inches(1.2))
    para(doc, ORG, size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    para(doc, f"UEN: {UEN}", size=8.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    para(doc, instrument.upper(), size=18, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    para(doc, "For", size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    if course_logo:
        p = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
        p.add_run().add_picture(str(course_logo), width=Inches(0.9))
    para(doc, COURSE_TITLE, size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=7)
    para(doc, f"TGS Ref No: {COURSE_CODE}", size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    para(doc, "Conducted by", size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    para(doc, ORG, size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    para(doc, f"UEN: {UEN}", size=8.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    para(doc, f"Version {VERSION}", size=10, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def toc_page(doc, document_label, entries):
    header = doc.add_paragraph()
    header.paragraph_format.space_after = Pt(18)
    r1 = header.add_run(document_label.upper())
    set_font(r1, size=10.5, bold=True, color=MUTED)
    r2 = header.add_run("  |  ")
    set_font(r2, size=10.5, bold=True, color=DARK)
    r3 = header.add_run(COURSE_TITLE)
    set_font(r3, size=10.5, color=MUTED)

    para(doc, "TABLE OF CONTENTS", size=16, bold=True, color=DARK, after=8)
    p = doc.paragraphs[-1]
    p_pr = p._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), RULE)
    p_bdr.append(bottom)
    for title, page, level in entries:
        row = doc.add_paragraph()
        row.paragraph_format.left_indent = Cm(0.55 * level)
        row.paragraph_format.space_after = Pt(7 if level == 0 else 5)
        row.paragraph_format.tab_stops.add_tab_stop(Inches(6.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        r1 = row.add_run(title)
        set_font(r1, size=11.2 if level == 0 else 10.4, bold=level == 0, color=DARK if level == 0 else MUTED)
        row.add_run("\t")
        r2 = row.add_run(str(page))
        set_font(r2, size=11.2 if level == 0 else 10.4, bold=level == 0, color=DARK)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def title_block(doc, title):
    para(doc, title, size=14, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, COURSE_TITLE, size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
    para(doc, f"{COURSE_CODE} | {VERSION}", size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)


def section_heading(doc, text):
    p = para(doc, text, size=12, bold=True, color=BLUE, before=8, after=5)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), RULE)
    p_bdr.append(bottom)


def trainee_info(doc):
    section_heading(doc, "A: Trainee Information")
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    rows = [("Name", ""), ("Last 3 NRIC Digits + Alphabet", ""), ("Date", "")]
    for row, (left, right) in zip(table.rows, rows):
        cell_text(row.cells[0], left, bold=True)
        cell_text(row.cells[1], right)
        shade_cell(row.cells[0], LIGHT)
    para(doc, "", after=4)


def instructions(doc, practical=False):
    section_heading(doc, "B: Instructions to Candidate")
    duration = "90 minutes" if practical else "60 minutes"
    items = [
        "This is an individual, open-book assessment.",
        f"You have {duration} to complete this assessment.",
        "Answer all questions or complete all tasks.",
        "Use your own words and show enough detail for the assessor to evaluate your understanding.",
        "You may refer to the course slides, learner guide and completed lab work.",
        "Submit the completed assessment through the LMS/TMS portal as directed by the assessor.",
    ]
    if practical:
        items.extend([
            "Your practical answers must be based on the lab sequence completed in class.",
            "Attach screenshots, diagrams, tables or notes where requested.",
        ])
    add_bullets(doc, items, size=9.6)


def grading_criteria(doc, practical=False):
    section_heading(doc, "C: Grading Criteria")
    if practical:
        standard = (
            "Complete all 10 practical tasks and provide the required lab-aligned evidence. "
            "Each task must meet its stated validation checkpoint."
        )
    else:
        standard = (
            "Answer all 5 short-answer questions and demonstrate the underpinning knowledge "
            "required across the 10 AI-102 learning outcomes."
        )
    para(doc, standard, size=9.6, after=4)
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Assessment decision", "Competent (C) / Not Yet Competent (NYC)"),
        ("Reassessment", "Any unmet criterion must be corrected or reassessed in accordance with assessment policy."),
    ]
    for row, (label, value) in zip(table.rows, rows):
        cell_text(row.cells[0], label, bold=True, size=9.2)
        cell_text(row.cells[1], value, size=9.2)
        shade_cell(row.cells[0], LIGHT)


def answer_box(doc, caption="Candidate answer / evidence"):
    table = doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    cell_text(table.rows[0].cells[0], caption, bold=True, color=MUTED, size=9.5)
    shade_cell(table.rows[0].cells[0], LIGHT)
    table.rows[1].cells[0].text = "\n\n\n\n\n"
    para(doc, "", after=3)


def official_use(doc):
    section_heading(doc, "For Official Use Only")
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    data = [
        ["Assessment Result", "C / NYC", "Assessor Name", ""],
        ["Assessor NRIC", "", "Date", ""],
        ["Assessor Signature", "", "Candidate Signature", ""],
        ["Remarks", "", "", ""],
    ]
    for r, row in enumerate(data):
        for c, value in enumerate(row):
            cell_text(table.rows[r].cells[c], value, bold=c in (0, 2), size=9.5)
            if c in (0, 2):
                shade_cell(table.rows[r].cells[c], LIGHT)


def alignment_table(doc, labs):
    section_heading(doc, "Assessment Alignment")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Area", "Slides", "Labs", "Assessment Coverage"]
    for i, h in enumerate(headers):
        cell_text(table.rows[0].cells[i], h, bold=True, color="FFFFFF")
        shade_cell(table.rows[0].cells[i], BLUE)
    for lab in labs:
        start = (lab["no"] - 1) * 4 + 1
        end = start + len(lab["questions"]) - 1
        row = table.add_row()
        row_data = [lab["short"], lab["slides"], f"Lab {lab['no']}", f"K{start}-K{end}, {lab['code']}"]
        for i, value in enumerate(row_data):
            cell_text(row.cells[i], value, size=8.6)


def written_questions(labs):
    return [
        {
            "code": "K1",
            "source": f"PPT slides {labs[0]['slides']} and {labs[1]['slides']}; Labs 1-2",
            "context": "An organization is planning a secure, production-ready Azure AI solution that handles customer data and potentially harmful content.",
            "question": "Explain how you would select and manage the Azure AI resources, secure access, monitor usage and cost, and apply responsible AI, content safety and governance controls.",
            "answer": [
                "Map business, data, integration, region, performance and cost requirements to the appropriate Azure AI services and deployment model.",
                "Protect keys, prefer managed identity and least privilege, separate environments, rotate secrets and secure network access.",
                "Monitor availability, latency, failures, consumption, cost, quality and safety signals with alerts and operational ownership.",
                "Apply risk assessment, content filters, prompt shields, privacy controls, human escalation, incident response and accountable governance.",
            ],
        },
        {
            "code": "K2",
            "source": f"PPT slides {labs[2]['slides']} and {labs[3]['slides']}; Labs 3-4",
            "context": "A customer-support assistant must answer from approved company documents and may use controlled tools to complete support actions.",
            "question": "Describe a grounded generative AI and agent solution, including RAG, prompt templates and evaluation, agent tools and permissions, orchestration, tracing, and human control of autonomous actions.",
            "answer": [
                "Use RAG to retrieve approved content, pass relevant context to the model and return grounded answers with traceable sources.",
                "Use prompt templates and prompt flow to standardize instructions, inputs, constraints and outputs, then evaluate quality, groundedness and safety.",
                "Give agents only approved tools and least-privilege permissions; orchestrate tool calls, memory, retrieval and any agent handoffs.",
                "Trace prompts, retrieval and tool actions, constrain high-impact actions, and require human approval or escalation where appropriate.",
            ],
        },
        {
            "code": "K3",
            "source": f"PPT slides {labs[4]['slides']} and {labs[5]['slides']}; Labs 5-6",
            "context": "A multilingual support workflow must understand images, scanned text, written messages and spoken interactions.",
            "question": "Explain how Azure vision, language, speech and translation capabilities could support this workflow, including when custom models, PII protection and human review are required.",
            "answer": [
                "Use image analysis, OCR, classification or object detection according to the visual requirement; train a custom vision model when prebuilt capabilities do not meet the domain need.",
                "Treat video and image data as sensitive and apply consent, retention, privacy, bias and access controls.",
                "Use language analysis for entities, sentiment, key phrases and PII detection before storing or sharing text.",
                "Use speech recognition, SSML-based synthesis and translation, with human review for sensitive, regulated, ambiguous or customer-facing outputs.",
            ],
        },
        {
            "code": "K4",
            "source": f"PPT slides {labs[6]['slides']} and {labs[7]['slides']}; Labs 7-8",
            "context": "Users need a conversational interface that recognizes support requests and searches a large multilingual knowledge base.",
            "question": "Design the language understanding, question answering and Azure AI Search components, covering intents, entities, training data, indexers, skillsets, vector search, semantic ranking, and backup or recovery.",
            "answer": [
                "Define representative intents, entities and varied utterances; train, test and improve the language project across expected wording and languages.",
                "Use custom question answering for curated answers and preserve project exports, labels, configuration and versions for recovery.",
                "Create a search index and data source, use an indexer to ingest content, and apply a skillset for OCR or other enrichment where required.",
                "Use vector search for semantic similarity and semantic ranking to improve relevance, then validate queries, security trimming and result quality.",
            ],
        },
        {
            "code": "K5",
            "source": f"PPT slides {labs[8]['slides']} and {labs[9]['slides']}; Labs 9-10",
            "context": "The final AI platform must extract information from varied documents and integrate the results into a governed, end-to-end customer-support architecture.",
            "question": "Explain how you would choose and validate Document Intelligence or Content Understanding models, integrate extracted content with search and generative AI, and operate the complete solution responsibly.",
            "answer": [
                "Use a prebuilt model for supported common forms, a custom model for domain-specific fields, and a composed model when multiple document types require routing.",
                "Use confidence scores and validation rules to accept results, trigger exceptions or route low-confidence and high-impact cases to human review.",
                "Use Content Understanding when the scenario spans documents, images, audio, video or text, and feed approved output into Azure AI Search and a grounded generative workflow.",
                "Document the service architecture, security, monitoring, cost, governance and cleanup plan; use RAG for current external knowledge and fine-tuning only when model behavior needs adaptation.",
            ],
        },
    ]


def build_wa_question(org_logo, course_logo, labs):
    questions = written_questions(labs)
    doc = base_doc()
    cover_page(doc, "Written Assessment", org_logo, course_logo)
    title_block(doc, "Written Assessment (Short Answer Questions)")
    trainee_info(doc)
    instructions(doc)
    grading_criteria(doc)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    section_heading(doc, "D: Written Questions")
    para(doc, "Answer all 5 questions. Each question is open-ended and aligned to the AI-102 PPT slides and labs.", italic=True, color=MUTED)
    for i, q in enumerate(questions, 1):
        para(doc, f"Question {i} ({q['code']})", size=11.5, bold=True, color=DARK, before=8, after=2)
        para(doc, f"Source: {q['source']}", size=9, color=MUTED, italic=True, after=2)
        para(doc, q["context"], size=10.5, color=MUTED, after=3)
        para(doc, q["question"], size=10.8, bold=True, after=5)
        answer_box(doc)
    official_use(doc)
    return doc


def build_wa_answer(org_logo, course_logo, labs):
    questions = written_questions(labs)
    doc = base_doc()
    cover_page(doc, "Answer to Written Assessment", org_logo, course_logo)
    title_block(doc, "Answer to Written Assessment (Marking Guide)")
    alignment_table(doc, labs)
    section_heading(doc, "Model Answers")
    para(doc, "Suggestive answers are not exhaustive. Award credit where the candidate demonstrates the same concept accurately.", italic=True, color=MUTED)
    for i, q in enumerate(questions, 1):
        para(doc, f"Question {i} ({q['code']}) - {q['source']}", size=11.5, bold=True, color=DARK, before=8, after=2)
        para(doc, q["question"], size=10.5, bold=True, after=3)
        add_bullets(doc, q["answer"], size=10.3)
    return doc


def build_pp_question(org_logo, course_logo, labs):
    doc = base_doc()
    cover_page(doc, "Practical Performance Assessment", org_logo, course_logo)
    title_block(doc, "Practical Performance Assessment")
    trainee_info(doc)
    instructions(doc, practical=True)
    grading_criteria(doc, practical=True)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    section_heading(doc, "D: Scenario")
    para(doc, "A company is designing a production-ready customer support AI platform with Azure AI services. Complete the practical tasks below using the same sequence, scenarios and deliverables practised in the labs.", size=10.8)
    section_heading(doc, "E: Practical Tasks")
    for i, lab in enumerate(labs, 1):
        para(doc, f"Task {i} ({lab['code']}) - {lab['short']}", size=11.5, bold=True, color=DARK, before=8, after=2)
        para(doc, f"Aligned to: Lab {lab['no']} and PPT slides {lab['slides']}", size=9, color=MUTED, italic=True, after=3)
        para(doc, lab["scenario"], size=10.8, after=5)
        para(doc, "Complete the lab sequence:", size=10.5, bold=True, after=2)
        add_bullets(doc, lab["steps"], style="List Number", size=10)
        para(doc, "Required evidence:", size=10.5, bold=True, after=2)
        add_bullets(doc, pp_evidence(lab), size=10)
        para(doc, f"Acceptance checkpoint: {lab['validation']}", size=10, color=MUTED, italic=True)
        answer_box(doc, "Candidate evidence / screenshots / notes")
    official_use(doc)
    return doc


def build_pp_answer(org_logo, course_logo, labs):
    doc = base_doc()
    cover_page(doc, "Answer to Practical Performance Assessment", org_logo, course_logo)
    title_block(doc, "Answer to Practical Performance Assessment (Marking Guide)")
    alignment_table(doc, labs)
    section_heading(doc, "Model Practical Evidence")
    para(doc, "The assessor should award credit where the candidate evidence follows the lab sequence and demonstrates the stated practical outcome.", italic=True, color=MUTED)
    for i, lab in enumerate(labs, 1):
        para(doc, f"Task {i} ({lab['code']}) - {lab['short']}", size=11.5, bold=True, color=DARK, before=8, after=2)
        para(doc, f"Aligned to: Lab {lab['no']} and PPT slides {lab['slides']}", size=9, color=MUTED, italic=True, after=2)
        para(doc, "Expected lab-aligned evidence:", size=10.5, bold=True, after=2)
        add_bullets(doc, pp_evidence(lab), size=10.3)
        para(doc, f"Validation standard: {lab['validation']}", size=10, color=MUTED, italic=True)
    return doc


def save_doc(doc, name):
    path = OUT / name
    doc.save(path)
    return path


def main():
    labs = load_labs()
    if len(labs) != 10:
        raise SystemExit(f"Expected 10 labs, found {len(labs)}")
    if len(written_questions(labs)) != 5:
        raise SystemExit("Expected exactly 5 written assessment questions")
    org_logo, course_logo = ensure_assets()
    files = [
        save_doc(build_wa_question(org_logo, course_logo, labs), f"WA (SAQ) - {COURSE_TITLE} - {VERSION}.docx"),
        save_doc(build_wa_answer(org_logo, course_logo, labs), f"Answer to WA (SAQ) - {COURSE_TITLE} - {VERSION}.docx"),
        save_doc(build_pp_question(org_logo, course_logo, labs), f"PP Assessment - {COURSE_TITLE} - {VERSION}.docx"),
        save_doc(build_pp_answer(org_logo, course_logo, labs), f"Answer to PP Assessment - {COURSE_TITLE} - {VERSION}.docx"),
    ]
    for file in files:
        print(file)


if __name__ == "__main__":
    main()
